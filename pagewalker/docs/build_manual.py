#!/usr/bin/env python3
"""
Build DOCX + PDF from PAGEWALKER_OPERATIONS_MANUAL.md

  pip install python-docx reportlab
  python3 pagewalker/docs/build_manual.py

Outputs next to this script:
  PAGEWALKER_OPERATIONS_MANUAL.docx
  PAGEWALKER_OPERATIONS_MANUAL.pdf
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

MD_NAME = "PAGEWALKER_OPERATIONS_MANUAL.md"
DOCX_NAME = "PAGEWALKER_OPERATIONS_MANUAL.docx"
PDF_NAME = "PAGEWALKER_OPERATIONS_MANUAL.pdf"


def _strip_md_bold(s: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*", r"\1", s)


def _parse_blocks(lines: list[str]) -> list[tuple[str, list[str]]]:
    """
    Yield (kind, lines) where kind is:
      h1, h2, h3, p, table, code, ascii
    """
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip("\n")
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i].rstrip("\n"))
                i += 1
            if i < n:
                i += 1
            yield ("code", block)
            continue
        if stripped.startswith("# "):
            yield ("h1", [stripped[2:].strip()])
            i += 1
            continue
        if stripped.startswith("## "):
            yield ("h2", [stripped[3:].strip()])
            i += 1
            continue
        if stripped.startswith("### "):
            yield ("h3", [stripped[4:].strip()])
            i += 1
            continue
        if "|" in stripped and stripped.startswith("|"):
            tbl = []
            while i < n:
                s = lines[i].strip()
                if not s or "|" not in s:
                    break
                if s.replace(" ", "") and all(c in "|-:" for c in s.replace(" ", "")):
                    i += 1
                    continue
                cells = [c.strip() for c in s.strip("|").split("|")]
                tbl.append(cells)
                i += 1
            if tbl:
                yield ("table", tbl)
            continue
        if stripped.startswith("┌") or stripped.startswith("└") or stripped.startswith("│") or stripped.startswith("─"):
            block = []
            while i < n and lines[i].strip():
                block.append(lines[i].rstrip("\n"))
                i += 1
            yield ("ascii", block)
            continue
        para = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if not nxt:
                break
            if nxt.startswith("#") or nxt.startswith("```") or nxt.startswith("|") or nxt.startswith("┌"):
                break
            para.append(nxt)
            i += 1
        yield ("p", para)


def build_docx(blocks: list[tuple[str, list[str]]], out: Path) -> None:
    from docx import Document

    doc = Document()
    doc.core_properties.title = "PageWalker Operations Manual"
    doc.core_properties.comments = "Generated from Markdown; edit MD and re-run build_manual.py"

    for kind, data in blocks:
        if kind == "h1":
            doc.add_heading(data[0], level=1)
        elif kind == "h2":
            doc.add_heading(data[0], level=2)
        elif kind == "h3":
            doc.add_heading(data[0], level=3)
        elif kind == "p":
            text = " ".join(_strip_md_bold(x) for x in data)
            p = doc.add_paragraph()
            p.add_run(text)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run("\n".join(data))
            run.font.name = "Courier New"
        elif kind == "ascii":
            p = doc.add_paragraph()
            run = p.add_run("\n".join(data))
            run.font.name = "Courier New"
        elif kind == "table":
            if not data:
                continue
            cols = max(len(r) for r in data)
            table = doc.add_table(rows=len(data), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(data):
                for ci in range(cols):
                    cell_text = row[ci] if ci < len(row) else ""
                    table.rows[ri].cells[ci].text = _strip_md_bold(cell_text)
        doc.add_paragraph()

    doc.save(out)


def build_pdf(blocks: list[tuple[str, list[str]]], out: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=12,
        textColor=colors.HexColor("#c2410c"),
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=14,
        spaceBefore=14,
        spaceAfter=8,
    )
    h3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontSize=12,
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=8,
    )
    mono = ParagraphStyle(
        "Mono",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=8,
        leading=10,
        leftIndent=12,
    )

    story: list = []
    for kind, data in blocks:
        if kind == "h1":
            story.append(Paragraph(escape(data[0]), h1))
        elif kind == "h2":
            story.append(Spacer(1, 0.1 * inch))
            story.append(Paragraph(escape(data[0]), h2))
        elif kind == "h3":
            story.append(Paragraph(escape(data[0]), h3))
        elif kind == "p":
            text = escape(" ".join(_strip_md_bold(x) for x in data))
            story.append(Paragraph(text.replace("\n", "<br/>"), body))
        elif kind in ("code", "ascii"):
            text = escape("\n".join(data))
            story.append(Paragraph(f"<font face='Courier'>{text.replace(chr(10), '<br/>')}</font>", mono))
        elif kind == "table" and data:
            t = Table([[escape(_strip_md_bold(c)) for c in row] for row in data])
            t.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#fff7ed")),
                        ("FONT", (0, 0), (-1, -1), "Helvetica", 8),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(Spacer(1, 0.05 * inch))
            story.append(t)
            story.append(Spacer(1, 0.1 * inch))
        story.append(Spacer(1, 0.06 * inch))

    doc = SimpleDocTemplate(
        str(out),
        pagesize=letter,
        rightMargin=inch * 0.75,
        leftMargin=inch * 0.75,
        topMargin=inch * 0.75,
        bottomMargin=inch * 0.75,
    )
    doc.build(story)


def main() -> int:
    here = Path(__file__).resolve().parent
    md_path = here / MD_NAME
    if not md_path.is_file():
        print(f"Missing {md_path}", file=sys.stderr)
        return 1
    text = md_path.read_text(encoding="utf-8")
    blocks = list(_parse_blocks(text.splitlines()))
    try:
        build_docx(blocks, here / DOCX_NAME)
    except ImportError:
        print("Install: pip install python-docx", file=sys.stderr)
        return 2
    try:
        build_pdf(blocks, here / PDF_NAME)
    except ImportError:
        print("Install: pip install reportlab", file=sys.stderr)
        return 3
    print(f"Wrote {here / DOCX_NAME}")
    print(f"Wrote {here / PDF_NAME}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
